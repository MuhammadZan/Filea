"""
PDF Converter Module
Supports PDF to Word and Word to PDF conversions
"""
import os
from pdf2docx import Converter as PDFConverter
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch


class DocumentConverter:
    """Handle PDF and Word document conversions"""
    
    SUPPORTED_FORMATS = ['pdf', 'docx']
    
    def __init__(self):
        """Initialize the document converter"""
        pass
    
    @staticmethod
    def is_supported(format_name):
        """Check if format is supported"""
        return format_name.lower() in DocumentConverter.SUPPORTED_FORMATS
    
    @staticmethod
    def pdf_to_word(pdf_path, docx_path):
        """
        Convert PDF to Word document
        
        Args:
            pdf_path (str): Path to input PDF file
            docx_path (str): Path to save Word document
        
        Returns:
            str: Path to converted file
        
        Raises:
            FileNotFoundError: If input file doesn't exist
            Exception: If conversion fails
        """
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")
        
        try:
            # Create converter instance
            cv = PDFConverter(pdf_path)
            
            # Convert PDF to DOCX
            cv.convert(docx_path, start=0, end=None)
            cv.close()
            
            return docx_path
            
        except Exception as e:
            raise Exception(f"PDF to Word conversion failed: {str(e)}")
    
    @staticmethod
    def word_to_pdf(docx_path, pdf_path):
        """
        Convert Word document to PDF
        
        Args:
            docx_path (str): Path to input Word document
            pdf_path (str): Path to save PDF file
        
        Returns:
            str: Path to converted file
        
        Raises:
            FileNotFoundError: If input file doesn't exist
            Exception: If conversion fails
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Word file not found: {docx_path}")
        
        try:
            # Load the Word document
            doc = Document(docx_path)
            
            # Create PDF
            pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=12,
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6,
            )
            
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    # Determine style based on paragraph style
                    if para.style.name.startswith('Heading'):
                        story.append(Paragraph(para.text, title_style))
                    else:
                        story.append(Paragraph(para.text, normal_style))
                    story.append(Spacer(1, 0.1 * inch))
            
            # Process tables
            for table in doc.tables:
                # Add a simple representation of tables
                for row in table.rows:
                    row_text = ' | '.join([cell.text for cell in row.cells])
                    if row_text.strip():
                        story.append(Paragraph(row_text, normal_style))
                story.append(Spacer(1, 0.2 * inch))
            
            # Build PDF
            pdf_doc.build(story)
            
            return pdf_path
            
        except Exception as e:
            raise Exception(f"Word to PDF conversion failed: {str(e)}")
    
    @staticmethod
    def convert(input_path, output_path, input_format, output_format):
        """
        Convert document between formats
        
        Args:
            input_path (str): Path to input file
            output_path (str): Path to save converted file
            input_format (str): Input format (pdf or docx)
            output_format (str): Output format (pdf or docx)
        
        Returns:
            str: Path to converted file
        
        Raises:
            ValueError: If conversion not supported
        """
        input_format = input_format.lower()
        output_format = output_format.lower()
        
        if not DocumentConverter.is_supported(input_format):
            raise ValueError(f"Unsupported input format: {input_format}")
        
        if not DocumentConverter.is_supported(output_format):
            raise ValueError(f"Unsupported output format: {output_format}")
        
        if input_format == output_format:
            raise ValueError("Input and output formats are the same")
        
        # Route to appropriate converter
        if input_format == 'pdf' and output_format == 'docx':
            return DocumentConverter.pdf_to_word(input_path, output_path)
        elif input_format == 'docx' and output_format == 'pdf':
            return DocumentConverter.word_to_pdf(input_path, output_path)
        else:
            raise ValueError(f"Conversion from {input_format} to {output_format} not supported")
    
    @staticmethod
    def get_document_info(file_path, format_type):
        """
        Get information about a document file
        
        Args:
            file_path (str): Path to document file
            format_type (str): Format type (pdf or docx)
        
        Returns:
            dict: Document information
        """
        info = {
            'format': format_type,
            'size': os.path.getsize(file_path) if os.path.exists(file_path) else 0
        }
        
        if format_type == 'docx':
            try:
                doc = Document(file_path)
                info['paragraphs'] = len(doc.paragraphs)
                info['tables'] = len(doc.tables)
            except:
                pass
        
        return info
